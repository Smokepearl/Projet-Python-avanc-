"""Génération du rapport Word avec python-docx.

Le rapport unique couvre les deux parties du projet :
    * Une page de titre : titre du livre, photo n° 1, auteur du livre,
      auteur du rapport.
    * Une section « Partie 1 » : résumé des données téléchargées depuis
      Internet (PokeAPI), agrégation calculée en SQL et graphique des données.
    * Une section « Partie 2 » : graphique de distribution des longueurs de
      paragraphes du premier chapitre + description (statistiques + source).

Les en-têtes utilisent différents styles avec des polices modifiées
(gras, italique), comme demandé.
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from . import charts
from .book_processor import BookData

ROOT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT_DIR = os.path.join(ROOT_DIR, "output")


def _add_custom_heading(doc: Document, text: str, size: int,
                        bold: bool, italic: bool, color: RGBColor) -> None:
    """Ajoute un en-tête avec une police personnalisée (gras/italique/couleur)."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.name = "Calibri"


def _add_part1_section(doc: Document, db) -> None:
    """Ajoute la section « Partie 1 » : données PokeAPI, agrégation SQL, graphique.

    :param db: instance ``database.Database`` contenant les données téléchargées.
    """
    heading = doc.add_heading("Partie 1 — Données téléchargées depuis Internet", level=1)
    for run in heading.runs:
        run.font.bold = True

    rows = db.fetch_all()
    if not rows:
        doc.add_paragraph(
            "Aucune donnée n'était présente dans la base au moment de la "
            "génération du rapport."
        )
        doc.add_page_break()
        return

    doc.add_paragraph(
        f"L'application a téléchargé {len(rows)} Pokémon au format JSON depuis "
        "l'API publique PokeAPI (https://pokeapi.co), puis a stocké un "
        "sous-ensemble des données (nom, taille, état, poids) dans une base "
        "SQLite. Le tableau ci-dessous présente une agrégation calculée par "
        "une requête SQL."
    )

    # Tableau d'agrégation SQL (taille en cm, poids en kg)
    doc.add_heading("Agrégation (requête SQL)", level=2)
    metrics = [("length", "Taille", "cm"), ("size", "Poids", "kg")]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for i, txt in enumerate(("Mesure", "Somme", "Moyenne", "Minimum", "Maximum")):
        run = table.rows[0].cells[i].paragraphs[0].add_run(txt)
        run.font.bold = True
    for column, label, unit in metrics:
        agg = db.aggregate(column)
        cells = table.add_row().cells
        cells[0].text = f"{label} ({unit})"
        cells[1].text = f"{agg['total']:.1f}"
        cells[2].text = f"{agg['moyenne']:.2f}"
        cells[3].text = f"{agg['minimum']:.1f}"
        cells[4].text = f"{agg['maximum']:.1f}"

    # Graphique des données de la base (poids)
    doc.add_heading("Graphique des données", level=2)
    chart_path = os.path.join(OUTPUT_DIR, "data_chart.png")
    fig = charts.build_db_figure(rows, column="size")
    charts.save_figure(fig, chart_path)
    pic_para = doc.add_paragraph()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.add_run().add_picture(chart_path, width=Inches(5.8))

    doc.add_page_break()


def generate_report(
    book: BookData,
    photo_path: str,
    report_author: str = "ELIDRISSI Hamza et GUY Michel",
    output_path: str | None = None,
    db=None,
) -> str:
    """Construit le document Word et l'enregistre. Renvoie le chemin du fichier.

    :param book: résultat de ``book_processor.analyze_book``.
    :param photo_path: chemin de « photo n° 1 » (image #1 + logo).
    :param db: instance ``database.Database`` (Partie 1). Si fournie, une
               section résumant les données téléchargées est ajoutée.
    """
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    if output_path is None:
        output_path = os.path.join(OUTPUT_DIR, "rapport.docx")

    doc = Document()
    navy = RGBColor(0x1B, 0x2A, 0x41)
    blue = RGBColor(0x2D, 0x6C, 0xDF)

    # ------------------------------------------------------------------ #
    # PAGE DE TITRE
    # ------------------------------------------------------------------ #
    doc.add_paragraph()  # petit espace en haut
    _add_custom_heading(doc, book.title, size=28, bold=True, italic=False, color=navy)
    _add_custom_heading(doc, f"par {book.author}", size=16, bold=False,
                        italic=True, color=blue)

    # Photo n° 1 (image #1 recadrée/redimensionnée + logo pivoté collé)
    if photo_path and os.path.exists(photo_path):
        pic_para = doc.add_paragraph()
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_para.add_run().add_picture(photo_path, width=Inches(3.2))

    doc.add_paragraph()
    _add_custom_heading(doc, "Rapport d'analyse", size=18, bold=True,
                        italic=False, color=navy)
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author_para.add_run(f"Auteur du rapport : {report_author}")
    run.font.italic = True
    run.font.size = Pt(13)

    doc.add_page_break()

    # ------------------------------------------------------------------ #
    # SECTION PARTIE 1 (données téléchargées + agrégation SQL + graphique)
    # ------------------------------------------------------------------ #
    if db is not None:
        _add_part1_section(doc, db)

    # ------------------------------------------------------------------ #
    # SECTION PARTIE 2 — page de graphique (distribution des paragraphes)
    # ------------------------------------------------------------------ #
    heading = doc.add_heading(
        "Partie 2 — Distribution des longueurs des paragraphes", level=1)
    for run in heading.runs:
        run.font.italic = True  # style d'en-tête modifié

    # Graphique enregistré sur disque puis inséré
    chart_path = os.path.join(OUTPUT_DIR, "distribution.png")
    fig = charts.build_paragraph_figure(book.distribution)
    charts.save_figure(fig, chart_path)

    chart_para = doc.add_paragraph()
    chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chart_para.add_run().add_picture(chart_path, width=Inches(5.8))

    # Description de l'intrigue / des données
    doc.add_heading("Description et statistiques", level=2)
    description = (
        f"Ce rapport analyse le premier chapitre de « {book.title} » de "
        f"{book.author}. Le graphique ci-dessus montre la distribution des "
        f"longueurs des paragraphes : chaque barre indique combien de "
        f"paragraphes possèdent un nombre de mots donné (arrondi à la dizaine "
        f"inférieure)."
    )
    doc.add_paragraph(description)

    stats = [
        ("Nombre de paragraphes", str(book.n_paragraphs)),
        ("Nombre total de mots (chapitre 1)", str(book.total_words)),
        ("Minimum de mots dans un paragraphe", str(book.min_words)),
        ("Maximum de mots dans un paragraphe", str(book.max_words)),
        ("Moyenne de mots par paragraphe", f"{book.avg_words:.1f}"),
        ("Source des données", book.source_url),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for label, value in stats:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        # Mettre le libellé en gras
        for run in cells[0].paragraphs[0].runs:
            run.font.bold = True

    try:
        doc.save(output_path)
    except PermissionError as exc:
        raise PermissionError(
            f"Impossible d'écrire « {output_path} » : le fichier est "
            "probablement ouvert dans Word. Fermez-le puis réessayez."
        ) from exc
    return output_path
